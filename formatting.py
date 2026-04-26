from fpdf import FPDF
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import os
import re
import tempfile
import matplotlib.pyplot as plt

def sanitize_text(text):
    """Clean text for FPDF compatibility (Standard fonts only support Latin-1)."""
    if not text:
        return ""
    # Map common unicode characters to latin-1 or similar
    chars = {
        '\u2013': '-', '\u2014': '-', 
        '\u2018': "'", '\u2019': "'", 
        '\u201c': '"', '\u201d': '"',
        '\u2022': '*', '\u2026': '...'
    }
    for u_char, r_char in chars.items():
        text = text.replace(u_char, r_char)
    return text.encode('latin-1', 'replace').decode('latin-1')

def generate_chart_images(dist):
    """Generates CO and Blooms chart images and returns paths to temporary files."""
    co_img = tempfile.NamedTemporaryFile(suffix='.png', delete=False).name
    bl_img = tempfile.NamedTemporaryFile(suffix='.png', delete=False).name
    
    # 1. CO Bar Chart
    co_data = dist.get('co', {})
    if co_data:
        fig, ax = plt.subplots(figsize=(6, 4))
        ax.bar(co_data.keys(), co_data.values(), color='#1f77b4')
        ax.set_title('Course Outcome Vs Mark Distribution', fontweight='bold')
        ax.set_ylabel('Marks')
        for i, v in enumerate(co_data.values()):
            ax.text(i, v + 0.2, str(v), ha='center')
        plt.tight_layout()
        plt.savefig(co_img)
        plt.close()
    else:
        co_img = None
        
    # 2. Blooms Pie Chart
    bl_data = dist.get('blooms', {})
    if bl_data:
        fig, ax = plt.subplots(figsize=(7, 5)) # Slightly larger for pie
        sorted_keys = sorted(bl_data.keys())
        active_labels = [k for k in sorted_keys if bl_data[k] > 0]
        active_vals = [bl_data[k] for k in active_labels]
        if active_vals:
            ax.pie(active_vals, 
                   labels=active_labels, 
                   autopct='%1.1f%%', 
                   colors=plt.cm.Paired.colors,
                   pctdistance=0.8,
                   labeldistance=1.1,
                   startangle=140)
            ax.set_title('Bloom\'s Level Vs Mark Distribution', fontweight='bold', pad=20)
            plt.tight_layout()
            plt.savefig(bl_img)
            plt.close()
        else:
            bl_img = None
    else:
        bl_img = None
        
    return co_img, bl_img

def create_pdf(paper):
    """Creates a PDF file from the generated question paper."""
    pdf = FPDF()
    pdf.add_page()
    
    # Header
    pdf.set_font("Helvetica", 'B', 16)
    pdf.multi_cell(0, 10, sanitize_text("DR.M.G.R. EDUCATIONAL AND RESEARCH INSTITUTE"), align='C')
    pdf.ln(2)
    pdf.set_font("Helvetica", 'B', 14)
    pdf.cell(0, 8, f"{paper.get('exam_type', 'CAT')} Examination", ln=True, align='C')
    pdf.ln(8)
    
    # Part A
    pdf.set_font("Helvetica", 'B', 12)
    pdf.cell(0, 8, "PART A - (10 x 1 = 10 Marks)", ln=True)
    pdf.set_font("Helvetica", '', 10)
    
    for q in paper.get('part_a', []):
        opts = q.get('options', {"A": "N/A", "B": "N/A", "C": "N/A", "D": "N/A"})
        pdf.set_font("Helvetica", 'B', 10)
        pdf.multi_cell(0, 6, sanitize_text(f"Q{q.get('number', '')}. {q.get('question', 'N/A')}"), 0, 'L')
        pdf.ln(1)
        pdf.set_font("Helvetica", '', 9)
        pdf.cell(10, 5, "")
        pdf.cell(0, 5, sanitize_text(f"A) {opts.get('A', 'N/A')}"), 0, 0, 'L')
        pdf.cell(40, 5, "")
        pdf.cell(0, 5, sanitize_text(f"B) {opts.get('B', 'N/A')}"), 0, 1, 'L')
        
        pdf.cell(10, 5, "")
        pdf.cell(0, 5, sanitize_text(f"C) {opts.get('C', 'N/A')}"), 0, 0, 'L')
        pdf.cell(40, 5, "")
        pdf.cell(0, 5, sanitize_text(f"D) {opts.get('D', 'N/A')}"), 0, 1, 'L')
        
        pdf.set_font("Helvetica", 'I', 7)
        pdf.cell(10, 4, sanitize_text(f"(Unit: {q.get('unit', 'N/A')}, CO: {q.get('co', 'N/A')}, Blooms: {q.get('blooms', 'N/A')})"), 0, 1, 'L')
        pdf.ln(3)
        
    pdf.ln(5)
    
    # Part B
    pdf.set_font("Helvetica", 'B', 12)
    if paper.get('exam_type') == 'CAT1':
        pdf.cell(0, 8, "PART B - (2 x 10 = 20 Marks)", ln=True)
    else:
        pdf.cell(0, 8, "PART B - (4 x 10 = 40 Marks)", ln=True)
    pdf.set_font("Helvetica", '', 10)
    
    for q in paper.get('part_b', []):
        pdf.set_font("Helvetica", 'B', 10)
        pdf.multi_cell(0, 6, sanitize_text(f"Q{q.get('number', '')}. {q.get('question', 'N/A')}"), 0, 'L')
        pdf.ln(1)
        pdf.set_font("Helvetica", 'I', 7)
        pdf.cell(0, 4, sanitize_text(f"(Unit: {q.get('unit', 'N/A')}, CO: {q.get('co', 'N/A')}, Blooms: {q.get('blooms', 'N/A')})"), 0, 0, 'L')
        pdf.set_font("Helvetica", 'B', 8)
        pdf.cell(0, 4, "10 MARKS", 0, 1, 'R')
        pdf.ln(8)
    
    # Add Charts to PDF
    dist = paper.get('distributions', {})
    if dist:
        try:
            co_img_path, bl_img_path = generate_chart_images(dist)
            if co_img_path and os.path.exists(co_img_path):
                pdf.add_page()
                pdf.set_font("Helvetica", 'B', 14)
                pdf.cell(0, 10, "Visual Representation of Mark Distribution", ln=True, align='C')
                pdf.ln(5)
                # Centering the image (A4 width is approx 210mm)
                pdf.image(co_img_path, x=15, w=180)
                os.unlink(co_img_path)
            
            if bl_img_path and os.path.exists(bl_img_path):
                pdf.ln(10)
                pdf.image(bl_img_path, x=15, w=180)
                os.unlink(bl_img_path)
        except Exception as e:
            print(f"PDF Chart Error: {e}")
        
    return bytes(pdf.output())

def create_word_doc(paper, template_path=None, output_file=None):
    """Creates a Word Document by filling a template dynamically."""
    if not template_path:
        target = "CAT 1 AI.docx" if paper['exam_type'] == "CAT1" else "CAT-II CD QPP.docx"
        template_path = target if os.path.exists(target) else None

    if template_path and os.path.exists(template_path):
        doc = docx.Document(template_path)
    else:
        doc = docx.Document()

    # 1. Fill Question Paper Table (Part A and Part B)
    question_table = None
    col_map = {"no": 0, "q": 1, "marks": 2, "co": 3, "bl": 4}
    header_row_idx = 0

    for table in doc.tables:
        found_table = False
        for r_idx, row in enumerate(table.rows[:15]):
            try:
                txt = " ".join([c.text.strip().upper() for c in row.cells])
                if "Q.NO" in txt or "QUESTION" in txt:
                    header_row_idx = r_idx
                    question_table = table
                    found_table = True
                    for c_idx, cell in enumerate(row.cells):
                        t = cell.text.strip().upper()
                        if "Q" in t and ("NO" in t or t == "Q.N" or t == "Q.NO"): col_map["no"] = c_idx
                        elif "QUESTION" in t: col_map["q"] = c_idx
                        elif "MARKS" in t: col_map["marks"] = c_idx
                        elif "CO" in t: col_map["co"] = c_idx
                        elif "BL" in t or "BLOOM" in t: col_map["bl"] = c_idx
                    break
            except:
                continue
        if found_table:
            break

    if question_table:
        def fill_questions(qs, start_row, q_table, c_map):
            for q in qs:
                num = str(q.get('number', ''))
                clean_num = "".join(filter(str.isdigit, num))
                if not clean_num: continue
                
                for r_idx in range(start_row, len(q_table.rows)):
                    row = q_table.rows[r_idx]
                    try:
                        cell0 = row.cells[c_map["no"]].text.strip()
                        clean_cell0 = "".join(filter(str.isdigit, cell0))
                        
                        if clean_cell0 == clean_num:
                            # Fill Question Text
                            q_text = q.get('question', 'N/A')
                            if 'options' in q and q.get('options'):
                                o = q['options']
                                q_text += f"\nA) {o.get('A','')}\nB) {o.get('B','')}\nC) {o.get('C','')}\nD) {o.get('D','')}"
                            
                            # Use paragraphs to preserve formatting if needed, but simple text replacement for now
                            row.cells[c_map["q"]].text = q_text
                            
                            # Fill Metadata
                            if c_map["marks"] < len(row.cells):
                                row.cells[c_map["marks"]].text = str(q.get('marks', '10' if int(clean_num) > 10 else '1'))
                            if c_map["co"] < len(row.cells):
                                row.cells[c_map["co"]].text = q.get('co', 'CO1')
                            if c_map["bl"] < len(row.cells):
                                row.cells[c_map["bl"]].text = q.get('blooms', 'L1')
                            break
                    except:
                        continue

        data_start = header_row_idx + 1
        fill_questions(paper.get('part_a', []), data_start, question_table, col_map)
        fill_questions(paper.get('part_b', []), data_start, question_table, col_map)

    # 2. Fill CO Definitions Table
    cos_defs = paper.get('cos_definitions', [])
    if cos_defs:
        co_dict = {c['id'].upper().replace(" ", ""): c['description'] for c in cos_defs}
        for table in doc.tables:
            if len(table.rows) > 4:
                first_row_txt = " ".join([c.text.strip().upper() for c in table.rows[0].cells])
                if any(k in first_row_txt for k in ["CO1", "CO", "COURSE OUTCOME"]):
                    for row in table.rows:
                        label = row.cells[0].text.strip().upper().replace(" ", "").replace("-", "")
                        if label in co_dict:
                            row.cells[1].text = co_dict[label]
                        elif "COURSEOUTCOME" in label:
                            digit = "".join(filter(str.isdigit, label))
                            if digit and f"CO{digit}" in co_dict:
                                row.cells[1].text = co_dict[f"CO{digit}"]

    # 3. Fill Marks Distribution Table
    dist = paper.get('distributions', {})
    if dist:
        for table in doc.tables:
            # Look for a row that contains CO1, CO2, or L1, L2 etc.
            co_header_row_idx = -1
            for r_idx, row in enumerate(table.rows):
                try:
                    row_txt = " ".join([c.text.strip().upper() for c in row.cells])
                    # Flexible check for CO1 and CO2 or L1 and L2
                    if (re.search(r"CO[\s\-]*1", row_txt) and re.search(r"CO[\s\-]*2", row_txt)) or \
                       (re.search(r"L[\s\-]*1", row_txt) and re.search(r"L[\s\-]*2", row_txt)):
                        co_header_row_idx = r_idx
                        break
                except:
                    continue
            
            if co_header_row_idx != -1:
                try:
                    headers = [c.text.strip().upper().replace(" ", "").replace("-", "") for c in table.rows[co_header_row_idx].cells]
                    
                    # The marks row is usually the one immediately following the CO labels
                    # We scan up to 2 rows after to find a row that looks like it should contain marks
                    for offset in [1, 2]:
                        if co_header_row_idx + offset < len(table.rows):
                            mark_row = table.rows[co_header_row_idx + offset]
                            row_txt = " ".join([c.text.strip().upper() for c in mark_row.cells])
                            
                            # Update CO marks
                            for co_id, marks in dist.get('co', {}).items():
                                cid = co_id.upper().replace(" ", "").replace("-", "")
                                if cid in headers:
                                    idx = headers.index(cid)
                                    if idx < len(mark_row.cells):
                                        mark_row.cells[idx].text = str(marks)
                            
                            # Update Bloom's marks
                            for bl_id, marks in dist.get('blooms', {}).items():
                                bid = bl_id.upper().replace(" ", "").replace("-", "")
                                if bid in headers:
                                    idx = headers.index(bid)
                                    if idx < len(mark_row.cells):
                                        mark_row.cells[idx].text = str(marks)
                            
                            # Also look for a "TOTAL" cell in this row or the header
                            if "TOTAL" in headers:
                                total_idx = headers.index("TOTAL")
                                total_marks = sum(dist.get('co', {}).values())
                                if total_idx < len(mark_row.cells):
                                    mark_row.cells[total_idx].text = str(total_marks)
                except Exception as e:
                    print(f"Distribution Table Error: {e}")
            
        # Add Charts after the distribution table
        try:
            co_img_path, bl_img_path = generate_chart_images(dist)
            
            # Add a heading for the charts
            doc.add_paragraph().add_run("\nVisual Representation of Mark Distribution").bold = True
            
            # Add images side by side if possible, or one after another
            if co_img_path and os.path.exists(co_img_path):
                doc.add_picture(co_img_path, width=Inches(3.5))
                os.unlink(co_img_path)
            
            if bl_img_path and os.path.exists(bl_img_path):
                doc.add_picture(bl_img_path, width=Inches(3.5))
                os.unlink(bl_img_path)
                
            doc.add_page_break()
        except Exception as e:
            print(f"Error adding charts to Word: {e}")

    # 4. Fill Answer Key Table (Now skips Table 0 and uses STRICT detection)
    all_qs = paper.get('part_a', []) + paper.get('part_b', [])
    if all_qs:
        for t_idx, table in enumerate(doc.tables):
            if t_idx == 0: continue # CRITICAL: Skip the main question paper table to prevent overwriting!
            
            ak_header_row = -1
            target_cols = {"no": -1, "ans": -1, "marks": -1}
            
            # Scan top 20 rows for Answer Key headers
            for r_idx in range(min(20, len(table.rows))):
                row_cells = table.rows[r_idx].cells
                row_txt = " ".join([c.text.strip().upper() for c in row_cells])
                
                # Rule out the main question paper table if it somehow got here
                if "QUESTION" in row_txt and "MARKS" in row_txt and "CO" in row_txt and "BL" in row_txt:
                    continue
                
                temp_cols = {"no": -1, "ans": -1, "marks": -1}
                for c_idx, cell in enumerate(row_cells):
                    t = cell.text.strip().upper()
                    if ("Q" in t and "NO" in t) or t == "Q.N" or t == "Q.NO": temp_cols["no"] = c_idx
                    elif "CONTENTS" in t or "ANSWER" in t: temp_cols["ans"] = c_idx
                    elif "MARKS" in t or "ALLOCATION" in t: temp_cols["marks"] = c_idx
                
                # Stricter check: Must have "CONTENTS" or "ALLOCATION" to be considered AK
                has_ak_marker = any(k in row_txt for k in ["CONTENTS", "ALLOCATION"])
                
                if temp_cols["no"] != -1 and temp_cols["ans"] != -1 and has_ak_marker:
                    ak_header_row = r_idx
                    target_cols = temp_cols
                    break
            
            if ak_header_row != -1:
                # We found an Answer Key table!
                for q in all_qs:
                    num_str = str(q.get('number', ''))
                    clean_num = "".join(filter(str.isdigit, num_str))
                    if not clean_num: continue
                    
                    for r_idx in range(ak_header_row + 1, len(table.rows)):
                        row = table.rows[r_idx]
                        if target_cols["no"] < len(row.cells):
                            cell0 = row.cells[target_cols["no"]].text.strip()
                            clean_cell0 = "".join(filter(str.isdigit, cell0))
                            
                            if clean_cell0 == clean_num:
                                ans_text = ""
                                if 'options' in q and q.get('options'):
                                    key = q.get('answer', 'A')
                                    opt_val = q['options'].get(key, '')
                                    ans_text = f"{key}) {opt_val}"
                                else:
                                    ans_text = q.get('answer', 'N/A')
                                
                                if target_cols["ans"] < len(row.cells):
                                    row.cells[target_cols["ans"]].text = ans_text
                                if target_cols["marks"] < len(row.cells):
                                    m_val = str(q.get('marks', '10' if int(clean_num) > 10 else '1'))
                                    row.cells[target_cols["marks"]].text = m_val
                                break

    # Final Save
    if output_file:
        doc.save(output_file)
        return output_file
    else:
        f_out = io.BytesIO()
        doc.save(f_out)
        return f_out.getvalue()
